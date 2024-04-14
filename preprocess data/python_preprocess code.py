
import pandas as pd # data manipulation
import numpy as np ## numerical calculation
from sqlalchemy import create_engine
from urllib.parse import quote 
from getpass import getpass
from sklearn.preprocessing import StandardScaler
from sklearn.impute import SimpleImputer
import matplotlib.pyplot as plt # data visualization
from sqlalchemy import create_engine # connect to SQL database
from feature_engine.outliers import Winsorizer
import seaborn as sns
import statsmodels.api as sm
import scipy.stats as stats
from sklearn.preprocessing import PowerTransformer
import scipy.stats as stats
import pylab
from scipy import stats
from sklearn.preprocessing import PowerTransformer
from scipy.stats import skew, kurtosis, mode

user_name = 'root'
database = 'kit_machine'
your_password = 'prince123'
engine = create_engine(f'mysql+pymysql://{user_name}:%s@localhost:3306/{database}' % quote(f'{your_password}'))

df= pd.read_csv(r"C:\Users\Prince Kumar Gupta\OneDrive\Documents\project_163 document\KIT.csv")

df.to_sql('automotive_kit', con = engine, if_exists = 'replace', chunksize = 1000, index = False)

data=df.drop(columns={'Customer Code', 'Customer Name','OEM', 'Item Description',
       'Product type', 'Item Code','Total'})

# Set the 'KIT ITEM' column as the index
data.set_index('KIT ITEM', inplace=True)

# Transpose the DataFrame to swap rows and columns
data = data.transpose()

# Identify duplicate column names
duplicate_columns = data.columns[data.columns.duplicated()].tolist()
duplicate_columns

# Aggregate duplicate columns
for col in duplicate_columns:
    data[col] = data[col].sum(axis=1)
    
    
data = data.loc[:, ~data.columns.duplicated()] ## remving the duplicate column 

## fill nan with zero ## 
data = data.fillna(0)

# Define threshold for proportion of zeros
zero_threshold = 0.20 # Columns with >80% zeros will be removed

# Calculate proportion of zeros in each column
zero_proportion = (data == 0).mean()

# Filter columns based on zero proportion threshold
high_zero_cols = zero_proportion[zero_proportion > zero_threshold].index

# Remove columns with high proportion of zeros
data1 = data.drop(columns=high_zero_cols)

## calculating the business moment #### 

def calculate_business_moments(dataframe, csv_filename):
    moments = {
        'Mean': dataframe.mean(),
        'Median': dataframe.median(),
        'Mode': dataframe.mode().iloc[0],  # Get the first mode
        'Minimum': dataframe.min(),
        'Maximum': dataframe.max(),
        'Skewness': dataframe.apply(skew),
        'Kurtosis': dataframe.apply(kurtosis),
        'Standard Deviation': dataframe.std(),
        'Variance': dataframe.var()
    }

    moments_df = pd.DataFrame(moments)
    moments_df=moments_df.transpose()
    moments_df.to_csv(csv_filename)
    
    return moments_df

# Call the function to calculate business moments and store them in a CSV file
data2=calculate_business_moments(data1, 'business_moments1.csv')

### imputation of missing value with median ######

data2=data1.replace(0, np.nan)

medians = data2.median()

data2=data2.fillna(medians)

data2.to_csv("preprocessed_data.csv")

import pandas as pd
from statsmodels.tsa.stattools import adfuller, kpss
from statsmodels.tsa.seasonal import seasonal_decompose
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from io import BytesIO

# Set Matplotlib style
plt.style.use('seaborn-darkgrid')

# Create a new Word document
doc = Document()

# Iterate over each column (time series) in the DataFrame
for column in data2.columns:
    doc.add_heading(f"Stationary tests for {column}", level=1)

    # ADF Test
    adf_result = adfuller(data2[column])
    doc.add_paragraph(f"ADF Test - p-value: {adf_result[1]}")
    if adf_result[1] <= 0.05:
        doc.add_paragraph("ADF Test: Series is stationary")
    else:
        doc.add_paragraph("ADF Test: Series is not stationary")

    # KPSS Test
    kpss_result = kpss(data2[column])
    doc.add_paragraph(f"KPSS Test - p-value: {kpss_result[1]}")
    if kpss_result[1] >= 0.05:
        doc.add_paragraph("KPSS Test: Series is stationary")
    else:
        doc.add_paragraph("KPSS Test: Series is not stationary")

    # Seasonality Trend Decomposition Plot
    decomposition = seasonal_decompose(data2[column], model='additive', period=12)  # Change period accordingly
    plt.figure(figsize=(10, 8))  # Adjust figure size for better readability
    plt.subplot(411)
    plt.plot(decomposition.observed)
    plt.ylabel('Original')
    plt.xticks(rotation=45)  # Rotate x-axis labels for better visibility
    plt.tight_layout()  # Adjust layout

    plt.subplot(412)
    plt.plot(decomposition.trend)
    plt.ylabel('Trend')
    plt.xticks(rotation=45)
    plt.tight_layout()

    plt.subplot(413)
    plt.plot(decomposition.seasonal)
    plt.ylabel('Seasonal')
    plt.xticks(rotation=45)
    plt.tight_layout()

    plt.subplot(414)
    plt.plot(decomposition.resid)
    plt.ylabel('Residual')
    plt.xticks(rotation=45)
    plt.tight_layout()

    # Adjust spacing between subplots
    plt.subplots_adjust(hspace=1.5)

    # Create a BytesIO object to temporarily store the plot
    tmp_img = BytesIO()
    plt.savefig(tmp_img, format='png')
    tmp_img.seek(0)

    # Add the plot to the Word document
    doc.add_picture(tmp_img, width=Inches(5))
    plt.close()

    # Add a page break for each column
    doc.add_page_break()

# Save the Word document
doc.save("stationarity_tests_output.docx")





